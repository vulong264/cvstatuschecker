"""
CV text extraction + Claude-powered structured parsing.

Pipeline:
  bytes (PDF/DOCX/TXT) → raw text → Claude → structured CandidateData
"""
import io
import json
import logging
import re
from dataclasses import dataclass, field
from typing import Optional

import anthropic
import pdfplumber
from docx import Document

from app.config import get_settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Structured output dataclass
# ---------------------------------------------------------------------------

@dataclass
class CandidateData:
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin_url: Optional[str] = None
    location: Optional[str] = None

    years_of_experience: Optional[float] = None
    current_title: Optional[str] = None
    current_company: Optional[str] = None

    main_skills: list[str] = field(default_factory=list)
    tech_stack: list[str] = field(default_factory=list)
    business_domains: list[str] = field(default_factory=list)

    education: list[dict] = field(default_factory=list)
    work_history: list[dict] = field(default_factory=list)

    cv_summary: Optional[str] = None
    raw_cv_text: Optional[str] = None


# ---------------------------------------------------------------------------
# Text extraction from raw bytes
# ---------------------------------------------------------------------------

def extract_text_from_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text(content: bytes, extension: str) -> str:
    """Extract plain text from CV bytes based on file extension."""
    ext = extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(content)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(content)
    elif ext == ".txt":
        return content.decode("utf-8", errors="replace")
    else:
        # Best-effort decode
        return content.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Claude parsing
# ---------------------------------------------------------------------------

EXTRACTION_PROMPT = """You are an expert HR analyst. Extract structured information from the CV text below.

Return a JSON object with EXACTLY these fields (use null for missing fields, empty arrays [] for missing lists):

{
  "full_name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "linkedin_url": "string or null",
  "location": "city, country or null",
  "years_of_experience": number_or_null,   // total professional years, e.g. 7.5
  "current_title": "string or null",
  "current_company": "string or null",
  "main_skills": ["skill1", "skill2"],     // top 10 most prominent skills
  "tech_stack": ["tech1", "tech2"],        // frameworks, DBs, cloud, tools
  "business_domains": ["domain1"],         // e.g. Fintech, E-commerce, Healthcare, SaaS
  "education": [
    {"degree": "BSc Computer Science", "institution": "University Name", "year": 2018}
  ],
  "work_history": [
    {"company": "Acme Corp", "role": "Senior Engineer", "years": 2.5, "description": "brief summary"}
  ],
  "cv_summary": "2-3 sentence professional summary of this candidate"
}

Rules:
- years_of_experience: calculate from earliest professional role to present; exclude internships if short
- main_skills: programming languages and core technical competencies first
- tech_stack: frameworks, libraries, databases, cloud platforms, DevOps tools
- business_domains: infer from project descriptions and industry context
- Be precise with emails and phone numbers (do not guess)
- Return ONLY valid JSON, no markdown fences, no extra text

CV TEXT:
---
{cv_text}
---"""


def parse_cv_with_claude(raw_text: str) -> CandidateData:
    """Send CV text to Claude and return a structured CandidateData object."""
    settings = get_settings()
    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    # Truncate extremely long CVs to stay within token limits
    truncated = raw_text[:15_000] if len(raw_text) > 15_000 else raw_text

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=2048,
        messages=[
            {
                "role": "user",
                "content": EXTRACTION_PROMPT.format(cv_text=truncated),
            }
        ],
    )

    response_text = message.content[0].text.strip()

    # Strip potential markdown code fences
    if response_text.startswith("```"):
        response_text = re.sub(r"^```(?:json)?\n?", "", response_text)
        response_text = re.sub(r"\n?```$", "", response_text)

    try:
        data = json.loads(response_text)
    except json.JSONDecodeError as exc:
        logger.error("Claude returned invalid JSON: %s\nResponse: %s", exc, response_text[:500])
        # Return minimal data with just raw text
        return CandidateData(raw_cv_text=raw_text)

    return CandidateData(
        full_name=data.get("full_name"),
        email=data.get("email"),
        phone=data.get("phone"),
        linkedin_url=data.get("linkedin_url"),
        location=data.get("location"),
        years_of_experience=data.get("years_of_experience"),
        current_title=data.get("current_title"),
        current_company=data.get("current_company"),
        main_skills=data.get("main_skills") or [],
        tech_stack=data.get("tech_stack") or [],
        business_domains=data.get("business_domains") or [],
        education=data.get("education") or [],
        work_history=data.get("work_history") or [],
        cv_summary=data.get("cv_summary"),
        raw_cv_text=raw_text,
    )


# ---------------------------------------------------------------------------
# High-level pipeline
# ---------------------------------------------------------------------------

def process_cv(content: bytes, extension: str) -> CandidateData:
    """
    Full pipeline: bytes → raw text → Claude extraction → CandidateData.
    Raises on extraction failure; Claude failures return minimal CandidateData.
    """
    raw_text = extract_text(content, extension)
    if not raw_text.strip():
        logger.warning("No text could be extracted from CV")
        return CandidateData(raw_cv_text="")
    return parse_cv_with_claude(raw_text)
